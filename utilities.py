import google.generativeai as genai
import os
import PyPDF2 as pdf
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from io import BytesIO
from reportlab.lib.pagesizes import letter
from dotenv import load_dotenv
import subprocess
import os

load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


def get_gemini_repsonse(input):
    model = genai.GenerativeModel("gemini-pro")
    response = model.generate_content(input)
    return response.text


def input_pdf_text(uploaded_file):
    reader = pdf.PdfReader(uploaded_file)
    text = ""
    for page in range(len(reader.pages)):
        page = reader.pages[page]
        text += str(page.extract_text())
    return text


def input_pdf_text_static(file_path):
    with open(file_path, "rb") as file:
        reader = pdf.PdfReader(file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text


def clean_keywords_response(response_text):
    try:
        # print("Received response_text:", response_text)  # Print the entire response_text

        # Parse the JSON-like string
        response_text = response_text.replace("“", '"').replace(
            "”", '"'
        )  # Replace smart quotes with standard quotes
        response_data = json.loads(response_text)  # Convert to dictionary

        # Extract the keywords
        if "Keywords" in response_data:
            keywords_part = response_data["Keywords"].strip().strip('"')
            keywords = [keyword.strip() for keyword in keywords_part.split(",")]
            return ", ".join(keywords) if keywords else "None"
        else:
            return "The 'Keywords' key is missing in the response."
    except json.JSONDecodeError:
        return "Failed to parse JSON response. Ensure the response is valid JSON."
    except Exception as e:
        return f"An error occurred: {str(e)}"

def clean_points_response(response_text):
    try:
        # Remove the "{"Points:"" prefix and the closing "}" if present
        cleaned_text = response_text.strip().lstrip('{"Points:"').rstrip('}').strip()
        
        # Split the text into individual points using '|'
        points = [point.strip() for point in cleaned_text.split('|') if point.strip()]
        
        # Format points with bullets and new lines
        formatted_points = "\n".join(f"• {point}" for point in points)
        
        return formatted_points
    except Exception as e:
        return f"An error occurred: {str(e)}"

def create_word_doc(content, filename):
    doc = Document()
    doc.add_heading("Cover Letter", 0)  # Optional heading

    # Add content to the document with proper text wrapping
    p = doc.add_paragraph()
    p.style.font.name = "Helvetica"
    p.style.font.size = Pt(12)

    # Add content to the paragraph
    p.add_run(content)

    # Save the documents
    doc.save(filename)


def create_pdf_from_text(text, title="Cover Letter"):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=52,
        leftMargin=52,
        topMargin=36,
        bottomMargin=18,
        title="Cover Letter Nikhil",
        author="Nikhil Kulkarni",
    )  # Reduced top margin

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(name="Justify", alignment=TA_JUSTIFY, fontSize=12)
    )  # Increased font size
    styles["Title"].fontSize = 18  # Increase title font size

    content = []

    # Add title
    content.append(Paragraph(title, styles["Title"]))
    content.append(Spacer(1, 12))

    # Add content paragraphs
    for para in text.split("\n"):
        if para.strip() == "":
            content.append(Spacer(1, 12))
        else:
            content.append(Paragraph(para, styles["Justify"]))
        content.append(Spacer(1, 6))

    doc.build(content)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

def load_latex_code(file_path):
    with open(file_path, 'r') as file:
        return file.read()
    
def save_latex_code(file_path, code):
    with open(file_path, 'w') as file:
        file.write(code)



def latex_to_pdf(latex_file,output_dir):
    # Check if the file exists
    if not os.path.isfile(latex_file):
        raise FileNotFoundError(f"The file {latex_file} does not exist.")
    

    # Run pdflatex command
    try:

        result = subprocess.run(
            ['pdflatex', '-output-directory', output_dir, latex_file],  # Use only the file name
            stdout=subprocess.PIPE,  # Capture output
            stderr=subprocess.PIPE,  # Capture errors
            text=True 
        )
        

        # Check if PDF was created
        pdf_file = latex_file.replace('.tex', '.pdf')
        if os.path.isfile(pdf_file):
            print(f"PDF successfully created: {pdf_file}")
            cleanup_auxiliary_files(latex_file, output_dir)
        else:
            print("PDF creation failed.")
    
    except subprocess.CalledProcessError as e:
        print(f"Subprocess failed with error: {e}")

def cleanup_auxiliary_files(latex_file, output_dir):
    """Delete auxiliary files such as .aux, .log, .out."""
    file_base = os.path.splitext(os.path.basename(latex_file))[0]  # Base name without extension
    extensions = ['.aux', '.log', '.out', '.toc','.tex']  # Add more extensions if needed

    for ext in extensions:
        aux_file = os.path.join(output_dir, file_base + ext)
        if os.path.isfile(aux_file):
            try:
                os.remove(aux_file)
            except Exception as e:
                print(f"Error deleting {aux_file}: {e}")

def extract_company_name(response):
    # Split the string by colon and take the first part (company name)
    if ":" in response:
        company_name = response.split(":")[0].strip()
        return company_name
    return None
